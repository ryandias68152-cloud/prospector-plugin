#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera o contrato .docx TRAVADO (somente leitura, com regiões editáveis para o cliente).
Uso: python3 gerar-docx.py dados.json saida.docx
dados.json: mesmas chaves do contrato-template.html (NOME_CLIENTE, VALOR, ... ) +
  "MANUTENCAO": true/false e "VALOR_MANUTENCAO" quando houver.
Campos que o CLIENTE preenche (regiões editáveis): CPF_CNPJ_CLIENTE, ENDERECO_CLIENTE (se vierem
como "(preencher)"), data e assinatura do contratante."""
import json, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

d = json.load(open(sys.argv[1], encoding='utf-8'))
PID = [100]

def par(doc, texto='', bold=False, center=False, size=11, antes=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(6)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if texto:
        r = p.add_run(texto); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Georgia'
    return p

def run(p, texto, bold=False, size=11):
    r = p.add_run(texto); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Georgia'
    return r

def editavel(p, texto):
    """Insere um trecho que o cliente PODE editar (permStart/permEnd, grupo everyone)."""
    PID[0] += 1; pid = str(PID[0])
    ps = OxmlElement('w:permStart'); ps.set(qn('w:id'), pid); ps.set(qn('w:edGrp'), 'everyone')
    p._p.append(ps)
    r = run(p, texto); r.font.highlight_color = 7  # amarelo: mostra onde preencher
    pe = OxmlElement('w:permEnd'); pe.set(qn('w:id'), pid)
    p._p.append(pe)

def campo(p, valor, rotulo):
    """Se o valor veio '(preencher)', vira região editável; senão texto normal."""
    if 'preencher' in (valor or '').lower() or not valor:
        editavel(p, ' [' + rotulo + ': preencher aqui] ')
    else:
        run(p, valor)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2); s.left_margin = s.right_margin = Cm(2.2)

par(doc, 'CONTRATO DE PRESTAÇÃO DE SERVIÇOS', bold=True, center=True, size=13)
par(doc, 'CRIAÇÃO E PUBLICAÇÃO DE PÁGINA NA INTERNET', bold=True, center=True, size=11)

p = par(doc); run(p, 'CONTRATANTE: ', bold=True); run(p, d['NOME_CLIENTE'] + ', ' + d['CPF_CNPJ_CLIENTE_LABEL'] + ' nº ')
campo(p, d.get('CPF_CNPJ_CLIENTE'), 'CPF/CNPJ'); run(p, ', com endereço em ')
campo(p, d.get('ENDERECO_CLIENTE'), 'endereço'); run(p, ', ' + d['CIDADE_UF_CLIENTE'] + '.')

p = par(doc); run(p, 'CONTRATADO(A): ', bold=True)
run(p, '%s, %s nº %s, com endereço em %s, %s.' % (d['NOME_PRESTADOR'], d['CPF_CNPJ_PRESTADOR_LABEL'], d['CPF_CNPJ_PRESTADOR'], d['ENDERECO_PRESTADOR'], d['CIDADE_UF_PRESTADOR']))

par(doc, 'As partes acima identificadas celebram o presente contrato de prestação de serviços, que se regerá pelas cláusulas seguintes.')

def clausula(n, titulo, texto):
    par(doc, 'Cláusula %sª — %s' % (n, titulo), bold=True, antes=12)
    par(doc, texto)

clausula(1, 'Do objeto', 'O presente contrato tem por objeto a criação de nova versão da página na internet do CONTRATANTE (%s), incluindo: redesign completo do layout com manutenção da identidade visual (logotipo, cores e imagens fornecidas), redação aprimorada do conteúdo existente, adaptação para dispositivos móveis e publicação da página no endereço %s.' % (d['URL_SITE_ANTIGO'], d['URL_PUBLICADA']))
clausula(2, 'Do valor e forma de pagamento', 'Pelos serviços descritos na Cláusula 1ª, o CONTRATANTE pagará ao CONTRATADO(A) o valor total de R$ %s (%s), na seguinte forma: %s.' % (d['VALOR'], d['VALOR_EXTENSO'], d['FORMA_PAGAMENTO']))
clausula(3, 'Do prazo de entrega', 'A página em sua versão final será entregue e publicada em até %s a contar da assinatura deste contrato e do fornecimento, pelo CONTRATANTE, dos materiais e aprovações necessários. Está incluída %s rodada(s) de ajustes de texto e imagens após a entrega.' % (d['PRAZO_ENTREGA'], d['RODADAS_AJUSTES']))
n = 4
if d.get('MANUTENCAO'):
    clausula(4, 'Da manutenção mensal', 'O CONTRATANTE contrata ainda o serviço de manutenção mensal da página (hospedagem, pequenas atualizações de texto/imagens e suporte), pelo valor de R$ %s mensais, com vigência a partir da publicação e renovação automática mensal.' % d['VALOR_MANUTENCAO'])
    n = 5
clausula(n, 'Do conteúdo e responsabilidades', 'O CONTRATANTE declara ser titular ou possuir autorização de uso de todos os textos, imagens, logotipo e informações fornecidos, responsabilizando-se pela veracidade das informações profissionais divulgadas. O CONTRATADO(A) compromete-se a não inserir na página informações não fornecidas ou não aprovadas pelo CONTRATANTE.')
clausula(n+1, 'Da hospedagem e domínio', d['TEXTO_HOSPEDAGEM'])
clausula(n+2, 'Da rescisão', 'Este contrato poderá ser rescindido por qualquer das partes mediante comunicação por escrito. Em caso de rescisão pelo CONTRATANTE após o início dos trabalhos, será devido o valor proporcional aos serviços já executados. Serviços de manutenção mensal, quando contratados, podem ser cancelados por qualquer parte com aviso prévio de 30 (trinta) dias.')
clausula(n+3, 'Do foro', 'Fica eleito o foro da comarca de %s para dirimir quaisquer controvérsias oriundas deste contrato.' % d['CIDADE_FORO'])

p = par(doc, antes=18); run(p, d['CIDADE_ASSINATURA'] + ', ')
editavel(p, ' [data] '); run(p, '.')

par(doc, '', antes=24)
p = par(doc, antes=18); run(p, '__________________________________________'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = par(doc, antes=0); run(p, d['NOME_CLIENTE'] + ' — Contratante  ', bold=True); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
editavel(p, ' [assine aqui] ')
p = par(doc, antes=18); run(p, '__________________________________________'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = par(doc, antes=0, center=True); run(p, d['NOME_PRESTADOR'] + ' — Contratado(a)', bold=True)

par(doc, 'Este documento é uma minuta base gerada automaticamente. Recomenda-se a revisão por profissional jurídico antes da assinatura. Gerado pelo plugin Prospector de Sites.', size=8, antes=20)

# PROTEÇÃO: somente leitura, exceto as regiões permitidas acima
dp = OxmlElement('w:documentProtection')
dp.set(qn('w:edit'), 'readOnly'); dp.set(qn('w:enforcement'), '1')
doc.settings.element.append(dp)

doc.save(sys.argv[2])
print('docx gerado:', sys.argv[2])
